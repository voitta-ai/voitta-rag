"""Parser for Microsoft Word documents (.docx)."""

from pathlib import Path

from .base import BaseParser, ParserResult


class DocxParser(BaseParser):
    """Parser for .docx files using python-docx."""

    extensions = [".docx"]

    def parse(self, file_path: Path) -> ParserResult:
        """Parse a Word document to Markdown."""
        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError:
            return ParserResult.failure("python-docx not installed. Run: pip install python-docx")

        try:
            doc = Document(file_path)
        except PackageNotFoundError:
            return ParserResult.failure(f"Invalid or corrupted DOCX file: {file_path}")
        except Exception as e:
            return ParserResult.failure(f"Failed to open DOCX file: {e}")

        lines = []
        metadata = {}

        # Extract core properties if available
        try:
            core_props = doc.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
        except Exception:
            pass

        # Add title as H1 if available
        if metadata.get("title"):
            lines.append(f"# {metadata['title']}")
            lines.append("")

        # Process paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check paragraph style for heading levels
            style_name = para.style.name.lower() if para.style else ""

            if "heading 1" in style_name:
                lines.append(f"# {text}")
            elif "heading 2" in style_name:
                lines.append(f"## {text}")
            elif "heading 3" in style_name:
                lines.append(f"### {text}")
            elif "heading 4" in style_name:
                lines.append(f"#### {text}")
            elif "heading 5" in style_name:
                lines.append(f"##### {text}")
            elif "heading 6" in style_name:
                lines.append(f"###### {text}")
            elif "list" in style_name or "bullet" in style_name:
                lines.append(f"- {text}")
            else:
                lines.append(text)

            lines.append("")

        # Process tables
        for table in doc.tables:
            table_md = self._table_to_markdown(table)
            if table_md:
                lines.append(table_md)
                lines.append("")

        content = "\n".join(lines).strip()
        return ParserResult(content=content, metadata=metadata)

    def _table_to_markdown(self, table) -> str:
        """Convert a Word table to Markdown format."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        # Build markdown table
        lines = []

        # Header row
        lines.append("| " + " | ".join(rows[0]) + " |")

        # Separator
        lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")

        # Data rows
        for row in rows[1:]:
            # Pad row if needed
            while len(row) < len(rows[0]):
                row.append("")
            lines.append("| " + " | ".join(row[: len(rows[0])]) + " |")

        return "\n".join(lines)
